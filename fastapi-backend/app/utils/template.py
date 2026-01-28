import os
from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docxtpl import DocxTemplate

from app.schemas import ResumeStructured


TEMPLATE_DIR = Path(__file__).resolve().parents[2] / "templates"
TEMPLATE_PATH = TEMPLATE_DIR / "resume_template.docx"


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _apply_font_doc(doc: Document, font_name: str = "Arial", font_size_pt: int = 10) -> None:
    styles = doc.styles
    if "Normal" in styles:
        normal = styles["Normal"]
        normal.font.name = font_name
        normal.font.size = Pt(font_size_pt)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = font_name
            r.font.size = Pt(font_size_pt)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = font_name
                        r.font.size = Pt(font_size_pt)
                        r._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _apply_hanging_indent_to_bullets(doc: Document, indent_pt: int = 18) -> None:
    left = Pt(indent_pt)
    first = Pt(-indent_pt)

    for p in doc.paragraphs:
        t = (p.text or "")
        if t.startswith("• "):
            pf = p.paragraph_format
            pf.left_indent = left
            pf.first_line_indent = first

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = (p.text or "")
                    if t.startswith("• "):
                        pf = p.paragraph_format
                        pf.left_indent = left
                        pf.first_line_indent = first


def ensure_template_exists() -> Path:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    _apply_font_doc(doc)

    tbl = doc.add_table(rows=16, cols=3)
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.columns[0].width = Inches(2.6)
    tbl.columns[1].width = Inches(2.2)
    tbl.columns[2].width = Inches(2.7)

    def set_row_height(row_idx: int, height_pt: int) -> None:
        row = tbl.rows[row_idx]
        row.height = Pt(height_pt)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    def merge_row_full(row_idx: int):
        c = tbl.cell(row_idx, 0)
        c.merge(tbl.cell(row_idx, 2))
        return c

    def merge_value_2cols(row_idx: int):
        return tbl.cell(row_idx, 1).merge(tbl.cell(row_idx, 2))

    hdr0 = merge_row_full(0)
    _set_cell_shading(hdr0, "9DC3E6")
    p0 = hdr0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("Candidate Submittal Cover Page")
    r0.bold = True

    cover_labels = [
        "Candidate Name:",
        "Phone Number:",
        "Email Address:",
        "Current location:",
        "Willing to Relocate?",
        "Former TCS employee /\ncontractor?\n(Please specify employment type and when)",
    ]
    cover_values = [
        "{{ name }}",
        "{{ phone }}",
        "{{ email }}",
        "{{ location }}",
        "{{ willing_to_relocate }}",
        "{{ former_tcs_employee_or_contractor }}",
    ]
    for i, (l, v) in enumerate(zip(cover_labels, cover_values), start=1):
        left = tbl.cell(i, 0)
        left.text = l
        _set_cell_shading(left, "D9EAF7")
        merge_value_2cols(i).text = v

    for i in range(1, 7):
        set_row_height(i, 30)

    hdr1 = merge_row_full(7)
    _set_cell_shading(hdr1, "9DC3E6")
    p1 = hdr1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("General Interview Availability")
    r1.bold = True

    instr = merge_row_full(8)
    _set_cell_shading(instr, "9DC3E6")
    pi = instr.paragraphs[0]
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri1 = pi.add_run(
        ' (ex. “Candidate is available every day during their lunch 12pm-1pm EST” OR “Candidate is available after 5pm EST everyday”)\n'
    )
    ri1.bold = True
    ri2 = pi.add_run("INCLUDE TIMEZONE")
    ri2.bold = True
    ri2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    left_av = tbl.cell(9, 0)
    left_av.text = "Candidate’s General\nAvailability:"
    _set_cell_shading(left_av, "D9EAF7")
    merge_value_2cols(9).text = "{{ interview_availability }}"

    hdr2 = merge_row_full(10)
    _set_cell_shading(hdr2, "9DC3E6")
    p2 = hdr2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Relevant Skills")
    r2.bold = True

    tbl.cell(11, 0).text = "Mandatory Skills\n(As listed in JD)"
    tbl.cell(11, 1).text = "# of Years\nExperience"
    tbl.cell(11, 2).text = "Candidate’s relevant hands-on\nexperience"
    for c in tbl.rows[11].cells:
        _set_cell_shading(c, "D9EAF7")
        if c.paragraphs:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx in range(1, 5):
        r = 11 + idx
        tbl.cell(r, 0).text = "{{ relevant_skill_" + str(idx) + "_skill }}"
        tbl.cell(r, 1).text = "{{ relevant_skill_" + str(idx) + "_years_required }}"
        tbl.cell(r, 2).text = "{{ relevant_skill_" + str(idx) + "_years_hands_on }}"

    for r in range(12, 16):
        set_row_height(r, 22)

    doc.add_paragraph("")
    h1 = doc.add_paragraph("SUMMARY")
    if h1.runs:
        h1.runs[0].bold = True
    doc.add_paragraph("{{ summary_bullets }}")

    doc.add_paragraph("")
    h2 = doc.add_paragraph("SKILLS")
    if h2.runs:
        h2.runs[0].bold = True
    doc.add_paragraph("{{ skills_bullets }}")

    doc.add_paragraph("")
    h3 = doc.add_paragraph("EXPERIENCE")
    if h3.runs:
        h3.runs[0].bold = True
    doc.add_paragraph("{{ experience_marker }}")

    doc.add_paragraph("")
    h4 = doc.add_paragraph("EDUCATION")
    if h4.runs:
        h4.runs[0].bold = True
    doc.add_paragraph("{{ education_marker }}")

    doc.add_paragraph("")
    h5 = doc.add_paragraph("PROJECTS")
    if h5.runs:
        h5.runs[0].bold = True
    doc.add_paragraph("{{ projects_marker }}")

    doc.add_paragraph("")
    h6 = doc.add_paragraph("CERTIFICATIONS")
    if h6.runs:
        h6.runs[0].bold = True
    doc.add_paragraph("{{ certifications_bullets }}")

    _apply_font_doc(doc)
    doc.save(str(TEMPLATE_PATH))
    return TEMPLATE_PATH


def _fmt_experience(r: ResumeStructured) -> str:
    lines: list[str] = []
    for item in r.experience:
        header = " - ".join([p for p in [item.title, item.company] if p]).strip()
        meta = " | ".join([p for p in [item.location, f"{item.start}–{item.end}".strip("–")] if p]).strip()
        if header:
            lines.append(f"• {header}")
        if meta:
            lines.append(f"• {meta}")
        for h in item.highlights:
            h = (h or "").strip()
            if h:
                lines.append(f"• {h}")
    return "\n".join(lines).strip()


def _render_experience_section(doc: Document, marker: str, r: ResumeStructured) -> None:
    marker_paragraphs = [p for p in doc.paragraphs if (p.text or "").strip() == marker]
    if not marker_paragraphs:
        return

    anchor = marker_paragraphs[0]

    def _tighten(p) -> None:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    spacer_after_heading = anchor.insert_paragraph_before("")
    _tighten(spacer_after_heading)

    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    usable_width = page_width - left_margin - right_margin

    for idx, item in enumerate(r.experience):
        if idx > 0:
            spacer_between = anchor.insert_paragraph_before("")
            _tighten(spacer_between)

        left_parts = [p for p in [item.company, item.location] if (p or "").strip()]
        left_text = " ".join([p.strip() for p in left_parts]).strip()

        end = (item.end or "").strip() or "Present"
        date_text = " to ".join([p for p in [item.start, end] if (p or "").strip()]).strip()

        header = anchor.insert_paragraph_before("")
        _tighten(header)
        pf = header.paragraph_format
        pf.tab_stops.add_tab_stop(usable_width, alignment=WD_TAB_ALIGNMENT.RIGHT)
        r1 = header.add_run(left_text)
        r1.bold = True
        header.add_run("\t")
        r2 = header.add_run(date_text)
        r2.bold = True

        if (item.title or "").strip():
            title_p = anchor.insert_paragraph_before(item.title.strip())
            _tighten(title_p)
            if title_p.runs:
                title_p.runs[0].bold = True

        highlights = [(h or "").strip() for h in (item.highlights or []) if (h or "").strip()]
        if highlights and re.match(r"^client\s*[-:]\s*", highlights[0], flags=re.IGNORECASE):
            client_p = anchor.insert_paragraph_before(highlights[0])
            _tighten(client_p)
            highlights = highlights[1:]

        for h in highlights:
            try:
                bp = anchor.insert_paragraph_before(h, style="List Bullet")
                _tighten(bp)
            except KeyError:
                p = anchor.insert_paragraph_before(h)
                p.text = f"• {h}"
                _tighten(p)

    parent = anchor._element.getparent()
    parent.remove(anchor._element)


def _render_projects_section(doc: Document, marker: str, r: ResumeStructured) -> None:
    marker_paragraphs = [p for p in doc.paragraphs if (p.text or "").strip() == marker]
    if not marker_paragraphs:
        return

    anchor = marker_paragraphs[0]

    def _tighten(p) -> None:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    spacer_after_heading = anchor.insert_paragraph_before("")
    _tighten(spacer_after_heading)

    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    usable_width = page_width - left_margin - right_margin

    for idx, item in enumerate(r.projects):
        name = (getattr(item, "name", "") or "").strip()
        desc = (getattr(item, "description", "") or "").strip()
        highlights = [(h or "").strip() for h in (getattr(item, "highlights", []) or []) if (h or "").strip()]

        if not (name or desc or highlights):
            continue

        if idx > 0:
            spacer_between = anchor.insert_paragraph_before("")
            _tighten(spacer_between)

        header_text = name or "Project"
        header = anchor.insert_paragraph_before("")
        _tighten(header)
        pf = header.paragraph_format
        pf.tab_stops.add_tab_stop(usable_width, alignment=WD_TAB_ALIGNMENT.RIGHT)
        r1 = header.add_run(header_text)
        r1.bold = True

        if desc:
            pdesc = anchor.insert_paragraph_before(desc)
            _tighten(pdesc)

        for h in highlights:
            try:
                bp = anchor.insert_paragraph_before(h, style="List Bullet")
                _tighten(bp)
            except KeyError:
                p = anchor.insert_paragraph_before(h)
                p.text = f"• {h}"
                _tighten(p)

    parent = anchor._element.getparent()
    parent.remove(anchor._element)


def _fmt_education(r: ResumeStructured) -> str:
    lines: list[str] = []
    for item in r.education:
        header = " - ".join([p for p in [item.degree, item.school] if p]).strip()
        meta = " | ".join([p for p in [item.location, f"{item.start}–{item.end}".strip("–")] if p]).strip()
        if header:
            lines.append(f"• {header}")
        if meta:
            lines.append(f"• {meta}")
    return "\n".join(lines).strip()


def _render_education_section(doc: Document, marker: str, r: ResumeStructured) -> None:
    marker_paragraphs = [p for p in doc.paragraphs if (p.text or "").strip() == marker]
    if not marker_paragraphs:
        return

    anchor = marker_paragraphs[0]

    def _tighten(p) -> None:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    spacer_after_heading = anchor.insert_paragraph_before("")
    _tighten(spacer_after_heading)

    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    usable_width = page_width - left_margin - right_margin

    for idx, item in enumerate(r.education):
        if idx > 0:
            spacer_between = anchor.insert_paragraph_before("")
            _tighten(spacer_between)

        degree = (item.degree or "").strip()
        date = (item.end or "").strip() or (item.start or "").strip()

        try:
            header = anchor.insert_paragraph_before("", style="List Bullet")
        except KeyError:
            header = anchor.insert_paragraph_before("")
        _tighten(header)
        pf = header.paragraph_format
        pf.left_indent = Pt(18)
        pf.first_line_indent = Pt(-18)
        pf.tab_stops.add_tab_stop(usable_width, alignment=WD_TAB_ALIGNMENT.RIGHT)
        r1 = header.add_run(degree)
        r1.bold = True
        header.add_run("\t")
        r2 = header.add_run(date)
        r2.bold = True

        school = (item.school or "").strip()
        loc = (item.location or "").strip()
        line2 = " － ".join([p for p in [school, loc] if p]).strip()
        if line2:
            p2 = anchor.insert_paragraph_before(line2)
            _tighten(p2)
            p2.paragraph_format.left_indent = Pt(18)
            p2.paragraph_format.first_line_indent = Pt(0)

    parent = anchor._element.getparent()
    parent.remove(anchor._element)


def _convert_bullet_text_to_word_bullets(doc: Document) -> None:
    def handle_paragraph(p) -> None:
        raw = (p.text or "")
        if not raw:
            return

        lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        if not lines:
            return

        if not all(ln.startswith("• ") for ln in lines):
            return

        bullet_texts = [ln[2:].strip() for ln in lines]

        for t in reversed(bullet_texts):
            try:
                p.insert_paragraph_before(t, style="List Bullet")
            except KeyError:
                p.insert_paragraph_before(t)

        parent = p._element.getparent()
        parent.remove(p._element)

    for p in list(doc.paragraphs):
        handle_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    handle_paragraph(p)


def _remove_paragraph(p) -> None:
    try:
        parent = p._element.getparent()
        parent.remove(p._element)
    except Exception:
        return


def _prune_section_by_marker(doc: Document, marker: str) -> None:
    ps = doc.paragraphs
    for i, p in enumerate(ps):
        if (p.text or "").strip() != marker:
            continue

        # Template layout is: blank line, HEADING, MARKER
        # Remove marker, heading, and preceding blank line if present.
        _remove_paragraph(p)
        if i - 1 >= 0:
            _remove_paragraph(ps[i - 1])
        if i - 2 >= 0:
            _remove_paragraph(ps[i - 2])
        return


def _prune_section_by_heading_and_next(doc: Document, heading: str) -> None:
    ps = doc.paragraphs
    for i, p in enumerate(ps):
        if (p.text or "").strip() != heading:
            continue

        # Template layout is: blank line, HEADING, CONTENT-PARAGRAPH
        # Remove content paragraph, heading, and preceding blank line if present.
        if i + 1 < len(ps):
            _remove_paragraph(ps[i + 1])
        _remove_paragraph(p)
        if i - 1 >= 0:
            _remove_paragraph(ps[i - 1])
        return


def _to_bullets(text: str) -> str:
    t = (text or "").strip()
    if not t:
        return ""

    parts = [p.strip() for p in t.splitlines() if p.strip()]
    if len(parts) <= 1:
        parts = [p.strip() for p in re.split(r"\.[\s\n]+", t) if p.strip()]
    return "\n".join([f"• {p.rstrip('.')}" for p in parts])


def render_resume_docx(structured: ResumeStructured, out_path: str) -> str:
    template_path = ensure_template_exists()
    tpl = DocxTemplate(str(template_path))

    links_line = " | ".join([l for l in structured.links if (l or "").strip()])
    skills_list = [s.strip() for s in structured.skills if (s or "").strip()]
    skills_bullets = f"• {', '.join(skills_list)}" if skills_list else ""

    context = {
        "name": structured.name,
        "email": str(structured.email or ""),
        "phone": structured.phone,
        "location": structured.location,
        "title": structured.title,
        "summary_bullets": _to_bullets(structured.summary),
        "willing_to_relocate": structured.willing_to_relocate,
        "former_tcs_employee_or_contractor": structured.former_tcs_employee_or_contractor,
        "interview_availability": structured.interview_availability,
        "interview_timezone": structured.interview_timezone,
        "links_line": links_line,
        "skills_bullets": skills_bullets,
        "experience_marker": "__EXPERIENCE_BLOCK__",
        "education_marker": "__EDUCATION_BLOCK__",
        "projects_marker": "__PROJECTS_BLOCK__",
        "certifications_bullets": "",
    }

    cert_lines: list[str] = []
    for c in (structured.certifications or []):
        name = (getattr(c, "name", "") or "").strip()
        issuer = (getattr(c, "issuer", "") or "").strip()
        date = (getattr(c, "date", "") or "").strip()
        left = " - ".join([p for p in [name, issuer] if p]).strip()
        right = date
        line = " | ".join([p for p in [left, right] if p]).strip()
        if line:
            cert_lines.append(f"• {line}")
    context["certifications_bullets"] = "\n".join(cert_lines).strip()

    for i in range(1, 5):
        item = structured.relevant_skills[i - 1] if len(structured.relevant_skills) >= i else None
        context[f"relevant_skill_{i}_skill"] = (item.skill if item else "")
        context[f"relevant_skill_{i}_years_required"] = (item.years_required if item else "")
        context[f"relevant_skill_{i}_years_hands_on"] = (item.years_hands_on if item else "")

    tpl.render(context)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    tpl.save(out_path)

    doc = Document(out_path)

    # Remove empty sections (heading + spacer + placeholder) so headings are dynamic.
    if not (context.get("summary_bullets") or "").strip():
        _prune_section_by_heading_and_next(doc, "SUMMARY")
    if not (context.get("skills_bullets") or "").strip():
        _prune_section_by_heading_and_next(doc, "SKILLS")
    if not (structured.experience or []):
        _prune_section_by_marker(doc, "__EXPERIENCE_BLOCK__")
    if not (structured.education or []):
        _prune_section_by_marker(doc, "__EDUCATION_BLOCK__")
    if not (structured.projects or []):
        _prune_section_by_marker(doc, "__PROJECTS_BLOCK__")
    if not (context.get("certifications_bullets") or "").strip():
        _prune_section_by_heading_and_next(doc, "CERTIFICATIONS")

    _apply_font_doc(doc)
    if structured.experience:
        _render_experience_section(doc, "__EXPERIENCE_BLOCK__", structured)
    if structured.education:
        _render_education_section(doc, "__EDUCATION_BLOCK__", structured)
    if structured.projects:
        _render_projects_section(doc, "__PROJECTS_BLOCK__", structured)
    _convert_bullet_text_to_word_bullets(doc)
    _apply_hanging_indent_to_bullets(doc)
    doc.save(out_path)
    return out_path

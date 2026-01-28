from docx import Document


def extract_docx_text(path: str) -> str:
    doc = Document(path)
    parts: list[str] = []

    def _is_list_paragraph(p) -> bool:
        try:
            style_name = (p.style.name or "").lower() if getattr(p, "style", None) else ""
            if "list" in style_name or "bullet" in style_name or "number" in style_name:
                return True
        except Exception:
            pass

        # Fallback: detect numbering properties
        try:
            ppr = p._p.pPr
            return bool(ppr is not None and ppr.numPr is not None)
        except Exception:
            return False

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        if _is_list_paragraph(p):
            parts.append(f"• {t}")
        else:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)

    return "\n".join(parts).strip()
